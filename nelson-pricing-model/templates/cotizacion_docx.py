"""
Template de cotización AlegentAI en .docx — python-docx.
Colores corporativos: azul oscuro #1F3764, azul medio #2E4A8C, naranja #C05000, verde #1A7A3C, gris #606060.
Uso: copiar y modificar las secciones de tablas con datos reales del proyecto.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

AZUL_OSCURO = RGBColor(0x1F, 0x37, 0x64)
AZUL_MEDIO  = RGBColor(0x2E, 0x4A, 0x8C)
GRIS        = RGBColor(0x60, 0x60, 0x60)
VERDE       = RGBColor(0x1A, 0x7A, 0x3C)
NARANJA     = RGBColor(0xC0, 0x50, 0x00)
HEADER_BG   = "2E4A8C"   # hex string para shd fill

def make_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)
    return doc

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    if p.runs: p.runs[0].font.color.rgb = AZUL_OSCURO
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    if p.runs: p.runs[0].font.color.rgb = AZUL_MEDIO
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def tabla(doc, headers, rows):
    """Tabla con header fondo azul + texto blanco."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hr = t.rows[0]
    for i, h in enumerate(headers):
        hr.cells[i].text = h
        for p in hr.cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hr.cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), HEADER_BG)
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
    return t

# ── PORTADA ──────────────────────────────────────────────────────────
# doc = make_doc()
# tit = doc.add_paragraph(); tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
# r = tit.add_run("PROPUESTA COMERCIAL")
# r.bold = True; r.font.size = Pt(24); r.font.color.rgb = AZUL_OSCURO
# ...
# doc.save("/tmp/cotizacion_CLIENTE_MMYY.docx")
